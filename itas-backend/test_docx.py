import docx
from django.test import Client
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
import os, django

os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'itas.settings')
django.setup()

# Create a DOCX file
doc = docx.Document()
doc.add_paragraph("Name: Mohammad Abumarar")
doc.add_paragraph("Email: mo.abumarar@gmail.com")
doc.add_paragraph("Title: Software Engineer")
doc.add_paragraph("Skills:")
doc.add_paragraph("Python, Django, React, Leadership")
doc.save("test_cv.docx")

# Test upload
User = get_user_model()
user = User.objects.filter(role="PM").first()
client = Client()
client.force_login(user)

with open("test_cv.docx", "rb") as f:
    file = SimpleUploadedFile("test_cv.docx", f.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response = client.post('/api/employees/analyze/', {'file': file}, SERVER_NAME="localhost", SERVER_PORT="8000")
    print("STATUS:", response.status_code)
    try:
        print("RESPONSE:", response.json())
    except:
        print("RAW:", response.content)
