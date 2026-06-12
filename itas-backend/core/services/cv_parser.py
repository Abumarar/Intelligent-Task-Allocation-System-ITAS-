"""
CV Parsing Service using PyPDF2
Extracts text from uploaded PDF CV files.
"""

import io
import os
import re
from typing import Any, Dict, Optional, Union

import joblib
import PyPDF2
import spacy
from django.conf import settings
from django.core.files.uploadedfile import UploadedFile

from core.services.text_preprocessor import clean_text


class ModelLoader:
    """Singleton for loading the ML model once."""

    _instance = None
    _model = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super(ModelLoader, cls).__new__(cls)
        return cls._instance

    @classmethod
    def get_model(cls) -> Any:
        if cls._model is None:
            cls._load_model()
        return cls._model

    @classmethod
    def _load_model(cls):
        model_path = os.environ.get("RESUME_MODEL_PATH")
        if not model_path:
            candidates = [
                os.path.join(settings.BASE_DIR, "resume_classifier_model.pkl"),
                os.path.join(settings.BASE_DIR, "..", "resume_classifier_model.pkl"),
                os.path.join(settings.BASE_DIR, "ai_training", "models", "domain_predictor.pkl"),
                os.path.join(settings.BASE_DIR, "..", "ai_training", "models", "domain_predictor.pkl"),
            ]
            for candidate in candidates:
                if os.path.exists(candidate):
                    model_path = candidate
                    break

        if not model_path or not os.path.exists(model_path):
            print("CV_PARSER_ERROR: Model file not found.")
            return

        try:
            print(f"CV_PARSER_INFO: Loading model from {model_path}...")
            model_obj = joblib.load(model_path)
            if isinstance(model_obj, dict) and "model" in model_obj:
                cls._model = model_obj["model"]
            else:
                cls._model = model_obj
            print("CV_PARSER_INFO: Model loaded successfully.")
        except Exception as e:
            print(f"CV_PARSER_ERROR: Failed to load model: {e}")


class CVParser:
    """Service for parsing CV/Portfolio PDF files."""

    @staticmethod
    def extract_text_from_pdf(
        file: Union[UploadedFile, io.BytesIO, bytes],
    ) -> Optional[str]:
        """
        Extract text content from a PDF file.

        Args:
            file: Django UploadedFile, BytesIO, or bytes object

        Returns:
            Extracted text as string, or None if parsing fails
        """
        try:
            # Handle different file inputs
            if isinstance(file, str) and os.path.isfile(file):
                with open(file, 'rb') as f:
                    file_content = f.read()
                file_stream = io.BytesIO(file_content)
            elif isinstance(file, bytes):
                file_stream = io.BytesIO(file)
            elif hasattr(file, "read"):
                if hasattr(file, "seek"):
                    file.seek(0)
                file_content = file.read()
                file_stream = io.BytesIO(file_content)
            else:
                return None

            # Create a PDF file reader
            try:
                pdf_reader = PyPDF2.PdfReader(file_stream)
            except PyPDF2.errors.PdfReadError:
                print("CV_PARSER_ERROR: Invalid or corrupted PDF file.")
                return None

            # Check if encrypted with explicit error
            if pdf_reader.is_encrypted:
                try:
                    # Try extracting without password (sometimes works for restriction-only encryption)
                    pass
                except Exception:
                    print("CV_PARSER_ERROR: Encrypted PDF. Cannot parse.")
                    return None

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    print(f"CV_PARSER_WARNING: Error parsing page {page_num}: {e}")
                    continue

            extracted_text = "\n\n".join(text_parts)
            return extracted_text.strip() if extracted_text else None

        except Exception as e:
            print(f"CV_PARSER_ERROR: Critical error parsing PDF: {str(e)}")
            return None

    @staticmethod
    def extract_text_from_docx(
        file: Union[UploadedFile, io.BytesIO, bytes],
    ) -> Optional[str]:
        """
        Extract text content from a DOCX file.

        Args:
            file: Django UploadedFile, BytesIO, or bytes object

        Returns:
            Extracted text as string, or None if parsing fails
        """
        import docx

        try:
            # Read the file content
            if isinstance(file, str) and os.path.isfile(file):
                with open(file, 'rb') as f:
                    file_content = f.read()
                doc_file = io.BytesIO(file_content)
            elif isinstance(file, bytes):
                doc_file = io.BytesIO(file)
            elif hasattr(file, "read"):
                if hasattr(file, "seek"):
                    file.seek(0)
                file_content = file.read()
                doc_file = io.BytesIO(file_content)
            else:
                return None

            doc = docx.Document(doc_file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            return "\n".join(full_text).strip()

        except Exception as e:
            print(f"Error parsing DOCX: {str(e)}")
            return None

    @staticmethod
    def is_valid_pdf(file: Union[UploadedFile, io.BytesIO, bytes]) -> bool:
        """Check if the uploaded file is a valid PDF."""
        try:
            if isinstance(file, bytes):
                file_stream = io.BytesIO(file)
            elif hasattr(file, "read"):
                file_content = file.read()
                if hasattr(file, "seek"):
                    file.seek(0)  # Reset file pointer
                file_stream = io.BytesIO(file_content)
            else:
                return False

            pdf_reader = PyPDF2.PdfReader(file_stream)
            return len(pdf_reader.pages) > 0
        except Exception:
            return False

    @staticmethod
    def is_valid_docx(file: Union[UploadedFile, io.BytesIO, bytes]) -> bool:
        """Check if the uploaded file is a valid DOCX."""
        import docx

        try:
            if isinstance(file, bytes):
                doc_file = io.BytesIO(file)
            elif hasattr(file, "read"):
                file_content = file.read()
                if hasattr(file, "seek"):
                    file.seek(0)  # Reset file pointer
                doc_file = io.BytesIO(file_content)
            else:
                return False

            docx.Document(doc_file)
            return True
        except Exception:
            return False

    @staticmethod
    def extract_details(text: str) -> Dict[str, Optional[str]]:
        """
        Extract basic details (Name, Email, Role/Title) from CV text.
        This is a heuristic implementation using spaCy.
        """
        details = {"name": None, "email": None, "title": None, "role": None}

        # Helper to clean weird PDF spacing (e.g. "S o f t w a r e" -> "Software", "e ngineer" -> "engineer")
        def _clean_spacing(s: str) -> str:
            if not s:
                return ""

            # 1. Merge sequences of single chars: "M o h a m m a d" -> "Mohammad"
            if re.search(r"\b([A-Za-z]\s){2,}[A-Za-z]\b", s):
                s = re.sub(r"([A-Za-z])\s(?=[A-Za-z]\b)", r"\1", s)

            # 2. Merge isolated single characters at start of words: "e ngineer" -> "engineer"
            s = re.sub(r"\b([A-Za-z])\s([A-Za-z]{3,})\b", r"\1\2", s)

            return re.sub(r"\s+", " ", s).strip()

        try:
            # 1. Extract Email (Highest Priority)
            email_match = re.search(
                r"[\w\.-]+@[\w\.-]+\.\w+", text
            )

            if email_match:
                details["email"] = email_match.group(0).strip()
            
            # Clean text for other extractions
            lines = [l.strip() for l in text.split("\n") if l.strip()]

            # 2. Extract Name
            # Labeled field "Name:"
            name_match = re.search(r"Name:\s*(.+)", text, re.IGNORECASE)
            if name_match:
                details["name"] = _clean_spacing(name_match.group(1))

            if not details["name"] and lines:
                # First-line heuristic with title guard
                first_line = _clean_spacing(lines[0])
                if 2 <= len(first_line.split()) <= 4 and re.match(
                    r"^[A-Za-z\s\.\-]+$", first_line
                ):
                    upper_first = first_line.upper()
                    title_keywords = [
                        "ENGINEER", "DEVELOPER", "MANAGER", "ANALYST", "DESIGNER", 
                        "ARCHITECT", "SCIENTIST", "ADMINISTRATOR", "CONSULTANT", 
                        "SPECIALIST", "LEAD", "DIRECTOR", "OFFICER", "RESUME", 
                        "CURRICULUM VITAE", "CV", "PROFILE", "SUMMARY"
                    ]
                    
                    is_title = any(kw in upper_first for kw in title_keywords)
                    
                    if not is_title:
                        # Email domain check
                        if details["email"] and first_line.lower().replace(" ", "") in details["email"].lower():
                            pass
                        elif "@" not in first_line:
                            details["name"] = first_line.title()

            # Try to extract name using LLM
            if not details["name"]:
                llm_name = CVParser.extract_name_with_llm(text)
                if llm_name:
                    details["name"] = llm_name
                    print(f"CV_PARSER_DEBUG: LLM Extracted Name: {llm_name}")

            # spaCy PERSON fallback
            if not details["name"]:
                try:
                    import spacy
                    nlp = spacy.load("en_core_web_sm")
                    first_chunk = "\n".join(text.splitlines()[:3])[:200]

                    cleaned_chunk = (
                        _clean_spacing(first_chunk)
                        if re.search(r"\b([A-Za-z]\s){2,}", first_chunk)
                        else first_chunk
                    )

                    chunk_doc = nlp(cleaned_chunk)

                    for ent in chunk_doc.ents:
                        if ent.label_ == "PERSON":
                            clean_name = ent.text.strip().title()
                            is_valid_name = True
                            if "@" in clean_name:
                                is_valid_name = False
                            elif details["email"] and clean_name.lower().replace(" ", "") in details["email"].lower():
                                is_valid_name = False

                            # Check against blocklist using substring matching for robust filtering
                            invalid_name_words = {
                                "software", "network", "system", "licensing", "management", 
                                "administration", "technology", "information", "project", "data", 
                                "business", "analyst", "engineer", "developer", "manager", 
                                "server", "database", "admin", "lead", "director", "coordinator",
                                "summary", "professional", "experience", "education", "skills",
                                "certifications", "hardware", "infrastructure", "troubleshooting",
                                "quality", "assurance", "testing", "operations", "vendor"
                            }
                            
                            if is_valid_name:
                                for word in clean_name.lower().split():
                                    if word in invalid_name_words:
                                        is_valid_name = False
                                        break

                            if (
                                is_valid_name
                                and len(clean_name.split()) >= 2
                                and re.match(r"^[A-Za-z\s\.\-]+$", clean_name)
                                and clean_name.lower() != "name"
                            ):
                                details["name"] = clean_name
                                break
                except OSError:
                    print("CV_PARSER_WARNING: spaCy model 'en_core_web_sm' not found. Skipping detailed name extraction.")

            # 3. Extract Role/Title
            # Labeled field "Role:/Title:/Position:"
            role_match = re.search(
                r"(?:Role|Title|Position):\s*(.+)", text, re.IGNORECASE
            )
            if role_match:
                details["title"] = _clean_spacing(role_match.group(1))

            roles_db = [
                "Software Engineer", "Java Developer", "Python Developer", 
                "Full Stack Developer", "Frontend Developer", "Backend Developer", 
                "DevOps Engineer", "Data Scientist", "Project Manager", 
                "Business Analyst", "Product Manager", "QA Engineer", 
                "UI/UX Designer", "Scrum Master", "System Administrator", 
                "Database Administrator"
            ]

            # Second-line check
            if not details["title"] and lines and len(lines) > 1:
                second_line = _clean_spacing(lines[1])
                # Check against full roles_db using the fixed regex
                for role in roles_db:
                    pattern = re.compile(re.escape(role), re.IGNORECASE)
                    if pattern.search(second_line):
                        details["title"] = role
                        break
                
                # Fallback to original second line logic
                if not details["title"] and ("|" in second_line or "Software" in second_line or "Developer" in second_line or "Engineer" in second_line):
                    details["title"] = second_line

            # Try to extract the title using LLM (Generative AI)
            if not details["title"]:
                llm_title = CVParser.extract_title_with_llm(text)
                if llm_title:
                    details["title"] = llm_title
                    print(f"CV_PARSER_DEBUG: LLM Extracted Title: {llm_title}")

            # roles_db scan
            if not details["title"]:
                cleaned_text = _clean_spacing(text[:1000])
                for role in roles_db:
                    pattern = re.compile(re.escape(role), re.IGNORECASE)
                    match = pattern.search(cleaned_text)
                    if match:
                        start, _ = match.span()
                        prefix_match = re.search(
                            r"\b(Senior|Sr\.|Junior|Jr\.|Lead|Principal|Chief)\s+$",
                            cleaned_text[:start],
                            re.IGNORECASE,
                        )
                        if prefix_match:
                            details["title"] = f"{prefix_match.group(1)} {role}".title()
                        else:
                            details["title"] = role.title()
                        break

        except Exception as e:
            print(f"CV_PARSER_ERROR: Error extracting details: {e}")

        print(f"CV_PARSER_DEBUG: Extracted text length: {len(text)}")
        print(f"CV_PARSER_DEBUG: Extracted details: {details}")
        return details

    @staticmethod
    def extract_task_details(text: str) -> Dict[str, Any]:
        """
        Extract task details (Title, Description, Skills, Priority) from document text.
        """
        details = {
            "title": None,
            "description": None,
            "requiredSkills": [],
            "priority": "MEDIUM",
        }

        if not text:
            return details

        lines = [l.strip() for l in text.split("\n") if l.strip()]

        # 1. Extract Title
        for line in lines[:5]:
            if line.lower().startswith("title:"):
                details["title"] = line.split(":", 1)[1].strip()
                break
            if line.lower().startswith("role:") or line.lower().startswith("position:"):
                details["title"] = line.split(":", 1)[1].strip()
                break

        if not details["title"] and lines:
            if len(lines[0]) < 100:
                details["title"] = lines[0]

        # 2. Extract Description
        description_start = 0
        for i, line in enumerate(lines):
            if "description" in line.lower() and len(line) < 30:
                description_start = i + 1
                break

        if description_start > 0:
            details["description"] = "\n".join(lines[description_start:])
        else:
            details["description"] = text[:2000]

        # 3. Extract Priority
        urgent_keywords = ["urgent", "immediate", "critical", "high priority"]
        if any(w in text.lower() for w in urgent_keywords):
            details["priority"] = "HIGH"

        return details

    @staticmethod
    def predict_role_with_ai(text: str, min_confidence: float = 0.45) -> Optional[str]:
        """
        Predict role using the trained AI model.
        """
        try:
            model = ModelLoader.get_model()
            if not model:
                return None

            cleaned = clean_text(text)
            if not cleaned:
                return None

            print(
                f"CV_PARSER_DEBUG: Model loaded. Predicting for text length {len(cleaned)}"
            )

            confidence = None
            if hasattr(model, "predict_proba"):
                probs = model.predict_proba([cleaned])[0]
                best_idx = int(probs.argmax())
                confidence = float(probs[best_idx])
                prediction = model.classes_[best_idx]
                print(f"CV_PARSER_DEBUG: Prediction confidence: {confidence:.3f}")
            else:
                prediction = model.predict([cleaned])[0]

            if confidence is not None and confidence < min_confidence:
                print("CV_PARSER_DEBUG: Low confidence prediction, skipping role.")
                return None

            print(f"CV_PARSER_DEBUG: Raw prediction: {prediction}")

            role = prediction.replace("-", " ").title()
            print(f"CV_PARSER_DEBUG: Formatted Role: {role}")
            return role

        except Exception as e:
            print(f"CV_PARSER_ERROR: Error predicting role: {e}")
            return None

    @staticmethod
    def extract_name_with_llm(text: str) -> str | None:
        """
        Use Google Gemini to extract the candidate's full name from a CV.
        """
        import os
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            print("CV_PARSER_DEBUG: GEMINI_API_KEY not found. Skipping LLM name extraction.", flush=True)
            return None
            
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=api_key)
            
            prompt = (
                "You are an expert HR assistant. Extract the candidate's full name "
                "from the following CV text. "
                "Return ONLY the exact name as a string, without any titles, degrees (like PMP), "
                "or extra words. If you cannot find a clear name, return the word 'NONE'.\n\n"
                f"CV Text:\n{text[:4000]}"
            )
            print(f"CV_PARSER_DEBUG: Sending prompt to Gemini for name with length: {len(prompt)}", flush=True)
            
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                ),
            )
            
            result = response.text.strip()
            print(f"CV_PARSER_DEBUG: Gemini name returned: '{result}'", flush=True)
            if result.upper() == 'NONE' or not result:
                return None
                
            return result.title()
        except Exception as e:
            import traceback
            print(f"CV_PARSER_ERROR: LLM Name Extraction failed: {e}", flush=True)
            return None

    @staticmethod
    def extract_title_with_llm(text: str) -> str | None:
        """
        Use Google Gemini to extract the most recent job title from a CV.
        """
        import os
        api_key = os.environ.get("GEMINI_API_KEY")
        if not api_key:
            print("CV_PARSER_DEBUG: GEMINI_API_KEY not found. Skipping LLM title extraction.", flush=True)
            return None
            
        try:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=api_key)
            
            prompt = (
                "You are an expert HR assistant. Extract the single most recent job title "
                "of the applicant from the following CV text. "
                "Return ONLY the exact job title as a string, nothing else. "
                "If it's an anonymized CV with no job experience or no clear title, return the word 'NONE'.\n\n"
                f"CV Text:\n{text[:4000]}"
            )
            print(f"CV_PARSER_DEBUG: Sending prompt to Gemini with length: {len(prompt)}", flush=True)
            
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.0,
                ),
            )
            
            result = response.text.strip()
            print(f"CV_PARSER_DEBUG: Gemini returned: '{result}'", flush=True)
            if result.upper() == 'NONE' or not result:
                return None
                
            return result.title()
        except Exception as e:
            import traceback
            print(f"CV_PARSER_ERROR: LLM Title Extraction failed: {e}", flush=True)
            traceback.print_exc()
            return None
